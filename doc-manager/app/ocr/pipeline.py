"""
Textara OCR pipeline — hosted Gemini vision backend.

PDF bytes in, DOCX bytes out. Each page is rendered to an image and sent to a
Google Gemini vision model with a plain "transcribe this Arabic page"
instruction; the returned text is assembled into a right-to-left Word document.

Replaces the former Qari / Qwen2-VL model that ran on Modal — no GPU, no torch,
no model loading. All OCR logic that is provider-agnostic (page rendering, DOCX
assembly) lives here so a different provider can be dropped in behind
``ocr_page`` without touching the worker.
"""

import io
import logging
import time

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image

from app.core.config import settings

log = logging.getLogger("doc-worker.ocr")

OCR_PROMPT = (
    "You are an OCR engine for Arabic documents. Transcribe ALL text visible in "
    "this page image exactly as it appears. Preserve the original line breaks and "
    "right-to-left reading order. Do NOT translate, summarize, correct, or add "
    "anything. If a word is unreadable, transcribe your best guess. Output ONLY the "
    "raw transcribed text — no preamble, no commentary, no markdown fences."
)


# ── Page rendering ────────────────────────────────────────────────────────────

def render_page(page) -> Image.Image:
    """Render a PDF page to a PIL image at the configured DPI."""
    mat = fitz.Matrix(settings.OCR_DPI / 72, settings.OCR_DPI / 72)
    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def _png_bytes(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


# ── OCR (Gemini) ──────────────────────────────────────────────────────────────

def ocr_page(img: Image.Image) -> str:
    """Run Gemini OCR on a single page image, returning the transcribed text.

    Retries transient 503/overload responses with linear backoff.
    """
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not set")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    contents = [
        types.Part.from_bytes(data=_png_bytes(img), mime_type="image/png"),
        OCR_PROMPT,
    ]

    last_exc = None
    for attempt in range(settings.OCR_MAX_RETRIES):
        try:
            resp = client.models.generate_content(
                model=settings.GEMINI_MODEL, contents=contents
            )
            return (resp.text or "").strip()
        except Exception as e:  # noqa: BLE001 — retry only on transient overload
            last_exc = e
            msg = str(e)
            if "503" in msg or "UNAVAILABLE" in msg or "overloaded" in msg.lower():
                time.sleep(2 * (attempt + 1))
                continue
            raise
    raise last_exc


# ── DOCX builder (RTL) ────────────────────────────────────────────────────────

def _set_rtl_paragraph(para) -> None:
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build_docx(pages_text: list[str], out) -> None:
    """Build a right-to-left DOCX from a list of per-page text strings."""
    doc = Document()

    sectPr = doc.sections[0]._sectPr
    sectPr.append(OxmlElement("w:bidi"))

    for page_num, text in enumerate(pages_text, 1):
        if page_num > 1:
            sep = doc.add_paragraph(f"-- Page {page_num} --")
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in sep.runs:
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.font.size = Pt(9)

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            para = doc.add_paragraph(line)
            _set_rtl_paragraph(para)
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(12)

    doc.save(out)


# ── Orchestration ─────────────────────────────────────────────────────────────

def process_pdf(pdf_bytes: bytes) -> bytes:
    """Take raw PDF bytes, return raw DOCX bytes."""
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    n_pages = len(pdf_doc)
    log.info("OCR: %d page(s) via %s", n_pages, settings.GEMINI_MODEL)

    pages_text: list[str] = []
    for i, page in enumerate(pdf_doc, 1):
        img = render_page(page)
        t0 = time.time()
        text = ocr_page(img)
        log.info("  page %d/%d — %d chars in %.1fs", i, n_pages, len(text), time.time() - t0)
        pages_text.append(text)
    pdf_doc.close()

    buf = io.BytesIO()
    build_docx(pages_text, buf)
    return buf.getvalue()
