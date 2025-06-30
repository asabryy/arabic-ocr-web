from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from transformers import Qwen2VLForConditionalGeneration, AutoProcessor
from pdf2image import convert_from_bytes
from qwen_vl_utils import process_vision_info
from docx import Document
import concurrent.futures
import os
import torch
import logging
from datetime import datetime
import re

app = FastAPI()

# Setup CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

# Global model and processor load
model_name = "NAMAA-Space/Qari-OCR-0.2.2.1-VL-2B-Instruct"
logger.info("Loading model and processor...")
model = Qwen2VLForConditionalGeneration.from_pretrained(
    model_name,
    torch_dtype=torch.float16,
    device_map="auto"
)
processor = AutoProcessor.from_pretrained(model_name)
prompt = (
            "Extract all Arabic text exactly as it appears in the image, including all diacritical marks (tashkeel), "
            "line breaks, paragraphs, and structure. Do not summarize. Do not skip any line."
        )
logger.info("Model and processor loaded.")

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "output"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# OCR function
def ocr_page(image):
    messages = [
        {"role": "user", "content": [
            {"type": "image", "image": image},
            {"type": "text", "text": prompt}
        ]}
    ]
    text_prompt = processor.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    image_inputs, _ = process_vision_info(messages)
    inputs = processor(text=[text_prompt], images=image_inputs, return_tensors="pt", padding=True).to(model.device)
    output = model.generate(**inputs, max_new_tokens=2000)
    generated_tokens = output[:, inputs["input_ids"].shape[-1]:]
    result = processor.batch_decode(generated_tokens, skip_special_tokens=True)[0]
    return result.strip()

# Sanitize filename
def sanitize_filename(filename):
    name = re.sub(r"[^a-zA-Z0-9_-]", "_", os.path.splitext(filename)[0])
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{name}_{timestamp}"

# OCR API
@app.post("/ocr")
async def upload_pdf(file: UploadFile = File(...)):
    safe_name = sanitize_filename(file.filename)
    input_path = os.path.join(UPLOAD_DIR, f"{safe_name}.pdf")
    output_path = os.path.join(OUTPUT_DIR, f"{safe_name}.docx")

    logger.info(f"Received file: {file.filename} | Saving to {input_path}")

    # Save PDF
    with open(input_path, "wb") as f:
        f.write(await file.read())

    # Convert PDF to images
    with open(input_path, "rb") as f:
        pdf_bytes = f.read()
    logger.info("Converting PDF to images...")
    images = convert_from_bytes(pdf_bytes, dpi=200)
    logger.info(f"Converted {len(images)} pages to images")

    # OCR each page
    logger.info("Running OCR in parallel...")
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        texts = list(executor.map(ocr_page, images))

    # Save DOCX
    logger.info(f"Writing output DOCX to {output_path}")
    doc = Document()
    for page_num, text in enumerate(texts, 1):
        doc.add_heading(f"Page {page_num}", level=2)
        doc.add_paragraph(text)
    doc.save(output_path)
    if not os.path.exists(output_path):
        logger.error(f"Failed to save DOCX: {output_path}")
        raise RuntimeError("DOCX file was not saved!")

    logger.info("OCR complete")
    return {"message": "OCR complete", "docx_path": f"/download/{safe_name}.docx"}

# Serve DOCX files for download
@app.get("/download/{filename}")
def download_file(filename: str):
    file_path = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(file_path):
        logger.info(f"Serving file for download: {file_path}")
        return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename)
    else:
        logger.warning(f"File not found: {file_path} (cwd: {os.getcwd()}) | Files in output: {os.listdir(OUTPUT_DIR)}")
        return {"error": "File not found"}

@app.get("/")
def health():
    return {"message": "Arabic OCR backend is running"}