from docx import Document

def save_docx(texts, path):
    doc = Document()
    for i, page in enumerate(texts, 1):
        doc.add_heading(f"Page {i}", level=2)
        doc.add_paragraph(page)
    doc.save(path)