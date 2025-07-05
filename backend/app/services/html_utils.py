from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

def set_rtl(run):
    # Add <w:rtl> to the run's properties
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def write_html_to_docx(html: str, output_path: str):
    doc = Document()
    soup = BeautifulSoup(html, "html.parser")

    # Remove unsupported <i> tag
    for i_tag in soup.find_all("i"):
        i_tag.unwrap()

    def add_paragraph(text, bold=False, underline=False, style=None):
        para = doc.add_paragraph(style=style)
        run = para.add_run(text)
        run.font.size = Pt(14)
        run.bold = bold
        run.underline = underline

        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        para.paragraph_format.right_to_left = True
        set_rtl(run)  # Ensure the run is also RTL

    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p"]):
        text = tag.get_text(strip=True)

        bold = tag.find("b") is not None
        underline = tag.find("u") is not None

        if tag.name.startswith("h"):
            level = int(tag.name[1])
            add_paragraph(text, bold=True, style=f"Heading {min(level, 3)}")
        else:
            add_paragraph(text, bold=bold, underline=underline)

    doc.save(output_path)
