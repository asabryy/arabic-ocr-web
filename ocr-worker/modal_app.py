"""
Textara OCR Pipeline — Modal GPU endpoint
==========================================
Receives a PDF as raw bytes, returns a DOCX as raw bytes.

Pipeline per page:
  1. PyMuPDF renders each page as a PIL image at 100 DPI
  2. Full-page image sent to Qari OCR (Qwen2-VL 2B) in one shot
  3. Assemble python-docx with RTL styles

Model loading:
  - bf16 Qwen2-VL-2B-Instruct base (SDPA attention)
  - Qari LoRA applied then merged (removes adapter overhead)
  - torch.compile(reduce-overhead) for faster repeated inference

Deploy:  modal deploy ocr-worker/modal_app.py
Call:    OCRPipeline = modal.Cls.from_name("textara-ocr", "OCRPipeline")
         docx_bytes  = OCRPipeline().process_pdf.remote(pdf_bytes)
"""

import io

import modal

image = (
    modal.Image.debian_slim(python_version="3.11")
    .apt_install(
        "libmupdf-dev",
        "mupdf-tools",
    )
    .pip_install(
        "numpy<2",
        "torch==2.4.0",
        "torchvision==0.19.0",
        "transformers>=4.49.0",
        "qwen-vl-utils",
        "accelerate>=0.26.0",
        "peft",
        "pymupdf",
        "Pillow",
        "python-docx",
    )
)

app = modal.App("textara-ocr", image=image)

BASE_MODEL_ID = "Qwen/Qwen2-VL-2B-Instruct"
LORA_MODEL_ID = "NAMAA-Space/Qari-OCR-0.2.2.1-VL-2B-Instruct"

DPI        = 100
MAX_PIXELS = 1280 * 28 * 28   # ~1 M pixels — full A4 page fits without downscaling

OCR_PROMPT = "Read and transcribe every word of Arabic text visible in this image, preserving the original line breaks."


@app.cls(
    gpu="A10G",
    timeout=600,
    min_containers=0,
)
class OCRPipeline:
    @modal.enter()
    def load_models(self):
        import time
        import torch
        from transformers import Qwen2VLForConditionalGeneration, AutoProcessor
        from peft import PeftModel

        t0 = time.time()

        print(f"Loading {BASE_MODEL_ID} (bf16 base, SDPA)...")
        base = Qwen2VLForConditionalGeneration.from_pretrained(
            BASE_MODEL_ID,
            torch_dtype=torch.bfloat16,
            device_map="auto",
            attn_implementation="sdpa",
        ).eval()

        print(f"Applying Qari LoRA ({LORA_MODEL_ID}) + merging weights...")
        model = PeftModel.from_pretrained(base, LORA_MODEL_ID)
        model = model.merge_and_unload()
        model.generation_config.temperature = None
        model.generation_config.top_p = None
        model.generation_config.top_k = None
        model.eval()

        self.processor = AutoProcessor.from_pretrained(BASE_MODEL_ID, use_fast=True)

        print("Compiling model (reduce-overhead)...")
        self.model = torch.compile(model, mode="reduce-overhead", fullgraph=False)

        elapsed = time.time() - t0
        vram_gb = torch.cuda.memory_allocated() / 1e9
        print(f"Model ready in {elapsed:.1f}s | VRAM: {vram_gb:.2f} GB")

    @modal.method()
    def process_pdf(self, pdf_bytes: bytes) -> bytes:
        """Takes raw PDF bytes, returns raw DOCX bytes."""
        import fitz
        import torch
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from PIL import Image
        from qwen_vl_utils import process_vision_info

        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        n_pages = len(pdf_doc)
        print(f"PDF: {n_pages} page(s)")

        pages_text = []
        for i, page in enumerate(pdf_doc, 1):
            print(f"  Page {i}/{n_pages} — rendering...")
            mat = fitz.Matrix(DPI / 72, DPI / 72)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            print(f"  Page {i}/{n_pages} — OCR ({img.width}x{img.height}px)...")
            messages = [{
                "role": "user",
                "content": [
                    {"type": "image", "image": img, "max_pixels": MAX_PIXELS},
                    {"type": "text", "text": OCR_PROMPT},
                ],
            }]

            text_input = self.processor.apply_chat_template(
                messages, tokenize=False, add_generation_prompt=True
            )
            image_inputs, video_inputs = process_vision_info(messages)
            inputs = self.processor(
                text=[text_input],
                images=image_inputs,
                videos=video_inputs,
                padding=True,
                return_tensors="pt",
            ).to(self.model.device)

            with torch.inference_mode():
                ids = self.model.generate(
                    **inputs,
                    max_new_tokens=2048,
                    do_sample=False,
                    num_beams=1,
                )

            trimmed = ids[:, inputs["input_ids"].shape[1]:]
            text = self.processor.batch_decode(
                trimmed,
                skip_special_tokens=True,
                clean_up_tokenization_spaces=True,
            )[0].strip()

            print(f"  Page {i}/{n_pages} — {len(text)} chars extracted")
            pages_text.append(text)

        pdf_doc.close()

        # Build DOCX
        doc = Document()

        sectPr = doc.sections[0]._sectPr
        bidi_el = OxmlElement("w:bidi")
        sectPr.append(bidi_el)

        for page_num, text in enumerate(pages_text, 1):
            if page_num > 1:
                sep = doc.add_paragraph(f"-- Page {page_num} --")
                sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in sep.runs:
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run.font.size = Pt(9)

            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                    continue
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                pPr = para._p.get_or_add_pPr()
                bidi = OxmlElement("w:bidi")
                bidi.set(qn("w:val"), "1")
                pPr.append(bidi)

                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(12)

        buf = io.BytesIO()
        doc.save(buf)
        print("Done.")
        return buf.getvalue()
