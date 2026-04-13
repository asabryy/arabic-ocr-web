"""
Textara OCR Pipeline — Qari (Qwen2-VL 2B)
==========================================
Core OCR logic shared by local_app.py (local dev) and modal_app.py (prod).

Each page is rendered as a full-page image and passed to Qari in one shot.
150 DPI gives ~1241x1754px for A4 — enough detail for the model to read
the full page without hitting EOS early.

CLI usage:
    python pipeline.py input.pdf -o output.docx
"""

import sys
import time
import argparse
from pathlib import Path

import fitz  # PyMuPDF
import torch
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ────────────────────────────────────────────────────────────────────

BASE_MODEL_ID = "Qwen/Qwen2-VL-2B-Instruct"
LORA_MODEL_ID = "NAMAA-Space/Qari-OCR-0.2.2.1-VL-2B-Instruct"

# 150 DPI -> A4 page ~= 1241x1754 px = ~2.2M pixels.
# MAX_PIXELS=4M lets the full page through without downscaling.
DPI        = 150
MAX_PIXELS = 1280 * 28 * 28 * 4

OCR_PROMPT = (
    "Below is the image of one page of a document, as well as some raw textual content "
    "that was previously extracted for it. Just return the plain text representation of "
    "this document as if you were reading it naturally. Do not hallucinate."
)

# ── Model ─────────────────────────────────────────────────────────────────────

def load_model():
    """Load Qari LoRA merged onto bf16 Qwen2-VL-2B base.

    Loads full-precision base, applies Qari LoRA, merges weights so inference
    uses native bf16 matmuls (no bnb dequantize overhead).
    SDPA in torch 2.4+ automatically uses Flash Attention 2 for bf16 on Ampere GPUs.
    Returns (model, processor, load_time_seconds).
    """
    from transformers import Qwen2VLForConditionalGeneration, AutoProcessor
    from peft import PeftModel

    t0 = time.time()

    print(f"Loading {BASE_MODEL_ID} (bf16 base, SDPA)...")
    base = Qwen2VLForConditionalGeneration.from_pretrained(
        BASE_MODEL_ID,
        torch_dtype=torch.bfloat16,
        device_map="auto",
        attn_implementation="sdpa",
    ).eval()

    print(f"Applying Qari LoRA ({LORA_MODEL_ID}) + merging weights...")
    model = PeftModel.from_pretrained(base, LORA_MODEL_ID)
    model = model.merge_and_unload()
    model.generation_config.temperature = None
    model.generation_config.top_p = None
    model.generation_config.top_k = None
    model.eval()

    processor = AutoProcessor.from_pretrained(BASE_MODEL_ID, use_fast=True)

    print("Compiling model (reduce-overhead)...")
    model = torch.compile(model, mode="reduce-overhead", fullgraph=False)

    elapsed = time.time() - t0
    vram_gb = torch.cuda.memory_allocated() / 1e9 if torch.cuda.is_available() else 0
    print(f"Model ready in {elapsed:.1f}s | VRAM: {vram_gb:.2f} GB")
    return model, processor, elapsed

# ── PDF ───────────────────────────────────────────────────────────────────────

def open_pdf(pdf_path: Path):
    """Open a PDF and return the fitz Document."""
    doc = fitz.open(str(pdf_path))
    print(f"PDF: {len(doc)} page(s)")
    return doc

# ── Page rendering ────────────────────────────────────────────────────────────

def render_page(page) -> Image.Image:
    """Render a PDF page to a PIL Image at DPI resolution."""
    mat = fitz.Matrix(DPI / 72, DPI / 72)
    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

# ── OCR ───────────────────────────────────────────────────────────────────────

def ocr_page(image: Image.Image, model, processor) -> tuple:
    """Run Qari OCR on a full-page PIL image.

    Returns (text, elapsed_seconds, vram_delta_bytes).
    """
    from qwen_vl_utils import process_vision_info

    vram_before = torch.cuda.memory_allocated() if torch.cuda.is_available() else 0
    t0 = time.time()

    messages = [{
        "role": "user",
        "content": [
            {"type": "image", "image": image, "max_pixels": MAX_PIXELS},
            {"type": "text",  "text": OCR_PROMPT},
        ],
    }]

    text_input = processor.apply_chat_template(
        messages, tokenize=False, add_generation_prompt=True
    )
    image_inputs, video_inputs = process_vision_info(messages)
    inputs = processor(
        text=[text_input],
        images=image_inputs,
        videos=video_inputs,
        padding=True,
        return_tensors="pt",
    ).to(model.device)

    with torch.inference_mode():
        ids = model.generate(
            **inputs,
            max_new_tokens=2048,
            do_sample=False,
            num_beams=1,
        )

    trimmed = ids[:, inputs["input_ids"].shape[1]:]
    text = processor.batch_decode(
        trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=True
    )[0].strip()

    elapsed    = time.time() - t0
    vram_delta = (torch.cuda.memory_allocated() if torch.cuda.is_available() else 0) - vram_before
    return text, elapsed, vram_delta

# ── DOCX builder ──────────────────────────────────────────────────────────────

def _set_rtl_paragraph(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build_docx(pages_text: list, out_path):
    """Build an RTL DOCX from a list of per-page text strings."""
    doc = Document()

    sectPr = doc.sections[0]._sectPr
    bidi_el = OxmlElement("w:bidi")
    sectPr.append(bidi_el)

    for page_num, text in enumerate(pages_text, 1):
        if page_num > 1:
            sep = doc.add_paragraph(f"-- Page {page_num} --")
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in sep.runs:
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.font.size = Pt(9)

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            para = doc.add_paragraph(line)
            _set_rtl_paragraph(para)
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(12)

    doc.save(out_path)
    if hasattr(out_path, "name"):
        print(f"Saved -> {out_path.name}")

# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Textara OCR Pipeline (Qari)")
    parser.add_argument("pdf",            help="Input PDF path")
    parser.add_argument("-o", "--output", help="Output DOCX path")
    args = parser.parse_args()

    pdf_path = Path(args.pdf).expanduser().resolve()
    if not pdf_path.exists():
        print(f"Error: {pdf_path} not found")
        sys.exit(1)
    out_path = Path(args.output) if args.output else pdf_path.with_suffix(".docx")

    print(f"\n{'='*60}")
    print(f"  Textara OCR Pipeline (Qari, full-page one-shot)")
    print(f"  Input : {pdf_path.name}")
    print(f"  Output: {out_path.name}")
    print(f"{'='*60}\n")

    print("[1/3] Loading model...")
    model, processor, load_time = load_model()

    print(f"\n[2/3] Processing pages...")
    pdf_doc = open_pdf(pdf_path)
    n = len(pdf_doc)

    pages_text = []
    for i, page in enumerate(pdf_doc, 1):
        print(f"  Page {i}/{n} — rendering...", end=" ", flush=True)
        img = render_page(page)
        print(f"({img.width}x{img.height}px) OCR...", end=" ", flush=True)
        text, elapsed, _ = ocr_page(img, model, processor)
        print(f"done ({elapsed:.1f}s, {len(text)} chars)")
        pages_text.append(text)
    pdf_doc.close()

    print(f"\n[3/3] Building DOCX...")
    build_docx(pages_text, out_path)
    print(f"\nComplete — {n} page(s) | model load {load_time:.1f}s\n")


if __name__ == "__main__":
    main()
